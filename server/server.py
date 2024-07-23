from flask import Flask, request, make_response, jsonify, send_file
from minio import Minio
import uuid
import json
import io
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
import requests
import hashlib

def getClient():
    minioClient = Minio('****',
        access_key='****',
        secret_key='****',
        secure=True
    )
    return minioClient

app = Flask(__name__)

# Подключение к минио с готовым бакетом
minioСlient = getClient()

def createWordFileIO(data):
    name_audio = data['filename']
    user_id = request.cookies.get('user_id')
    # Формируем имя файла для сохранения в бакете MinIO
    minio_filename = f"{user_id}%tx%{name_audio}"
    # Если объект существует, загружаем его и декодируем JSON
    dataMinio = minioСlient.get_object('interview-transcription-bucket', minio_filename)
    data_bytes = dataMinio.read()
    # Декодируем данные из JSON
    data_dict = json.loads(data_bytes.decode('utf-8'))
    # Создаем новый документ Word
    doc = Document()
    phrases = data_dict['phrases']
    size_font = data_dict['sizeFont']
        # Добавляем фразы в документ Word
    for phrase in phrases:
        input_time, speaker, phrase_text, phrase_bg_color = phrase
        # Добавляем абзац в документ
        paragraph = doc.add_paragraph()
        # Устанавливаем выравнивание абзаца (выравнивание по левому краю)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        # Создаем объекты для стилизации текста
        run = paragraph.add_run()
        # Добавляем текст в формате "inputTime \t speaker \t phraseText"
        run.text = f"{input_time}  ---  {speaker}  ---  {phrase_text}"
        # Устанавливаем размер шрифта (например, 12 points)
        font = run.font
        font.size = Pt(size_font - 5)

    temp_buffer = io.BytesIO()
    doc.save(temp_buffer)
    temp_buffer.seek(0)

    return temp_buffer

def createDataPhrase(result_neuro, filename):
    data = []
    for item in result_neuro:
        input_time = "{:02d}:{:02d}:{:02d}".format(int(item[0] // 3600), int((item[0] % 3600) // 60), int(item[0] % 60))
        new_element = [
            input_time, 
            item[3],
            item[2],   
            '#ffffff'  
        ]
        data.append(new_element)
    
    spikers = ['Спикер1', 'Спикер2']
    postData = {
        "phrases": data,
        "nameAudio": filename,
        "sizeFont": 14,
        "spikers": spikers
    }
    
    return postData

# Загрузка файла в хранилище Minio
@app.route('/upload/loadfile', methods=['POST'])
def upload_file():
    try:
        if 'myfile' not in request.files:
            return jsonify({'error': 'No file part'}), 400

        file = request.files['myfile']
        filename = request.form['filename']
        fileSize = int(request.form['fileSize'])
        numSpikers = request.form['numSpikers']

        if file:
            #Отправка файла в gen_neuro
            files = {'myfile': (filename, file.stream, file.mimetype)}
            data = {'numSpikers': numSpikers}
            response = requests.post('http://gen_neuro_server:5001/process_audio', files=files, data=data)

            if response.status_code == 200:
                #Создание данных транскрипции для минио
                resultsNeuro = response.json()
                postData = createDataPhrase(resultsNeuro, filename)

                # Получаем id из cookies
                user_id = request.cookies.get('user_id')

                #Сохраняем транскрипцию в минио
                minio_filename_trans = f"{user_id}%tx%{filename}"
                # Преобразуем данные в строку JSON
                data_json = json.dumps(postData)
                data_bytes = data_json.encode('utf-8')
                data_stream = io.BytesIO(data_bytes)
                # Загружаем данные в Minio
                minioСlient.put_object('interview-transcription-bucket', minio_filename_trans, data_stream, length=len(data_bytes))
                
                # Востанавливаем позицию курсора
                file.seek(0)
                # Формируем имя файла для сохранения в бакете MinIO
                minio_filename_audio = f"{user_id}%mp%{filename}"
                # Загружаем файл в бакет MinIO
                minioСlient.put_object('interview-transcription-bucket',
                                        minio_filename_audio,
                                        file,
                                        length=fileSize,
                                        )
                
                return jsonify({'success': True, 'message': 'File uploaded successfully'})
            else:
                return jsonify({'error': 'Error processing file on neuro server'}), response.status_code
        else:
            return jsonify({'success': False, 'error': 'File not provided'})
    
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

# Создание word транскрипции
@app.route('/upload/word', methods=['POST'])
def download_word_file():
    try:
        data = request.json 
        temp_buffer = createWordFileIO(data)

        # Отправляем файл в виде ответа
        return send_file(temp_buffer, as_attachment=True, download_name='output.docx',)

    except Exception as e:
        return jsonify({'error': str(e)}), 500

#Загрузка текста в минио
@app.route('/upload/savetext', methods=['POST'])
def upload_to_minio():
    try:
        data = request.json  
        name_audio = data['nameAudio']
        user_id = request.cookies.get('user_id')
        # Формируем имя файла для сохранения в бакете MinIO
        minio_filename = f"{user_id}%tx%{name_audio}"

        #Проверка существования файла.
        objects = minioСlient.list_objects('interview-transcription-bucket', prefix=minio_filename, recursive=True)
        first_object = next(objects, None)
        # Если фаил уже хранится, удаляет его
        if first_object is not None:
            minioСlient.remove_object('interview-transcription-bucket', minio_filename)
        
        # Преобразуем данные в строку JSON
        data_json = json.dumps(data)
        data_bytes = data_json.encode('utf-8')
        data_stream = io.BytesIO(data_bytes)
        # Загружаем данные в Minio
        minioСlient.put_object('interview-transcription-bucket', minio_filename, data_stream, length=len(data_bytes))
        
        return jsonify({'message': 'Данные успешно загружены в Minio'}), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 500
    
#Возврат текста клиенту
@app.route('/upload/gettext', methods=['POST'])
def load_to_client():
    try:
        data = request.json 
        name_audio = data['filename']
        user_id = request.cookies.get('user_id')
        minio_filename = f"{user_id}%tx%{name_audio}"

        # Если объект существует, загружаем его и декодируем JSON
        dataMinio = minioСlient.get_object('interview-transcription-bucket', minio_filename)
        # Читаем данные из объекта как байты
        data_bytes = dataMinio.read()
        # Преобразуем байты в строку JSON
        data_json = data_bytes.decode('utf-8')
        response_data = json.loads(data_json)

        return jsonify(response_data), 200
    except Exception as e:
       return jsonify({'error': str(e)}), 500

@app.route('/upload/getfile', methods=['POST'])
def get_audio():
    try:
        # Получаем название файла из тела запроса
        data = request.json
        filename = data['filename']
        if filename:
            user_id = request.cookies.get('user_id')
            minio_filename = f"{user_id}%mp%{filename}"
            # Получаем объект (аудиофайл) из бакета MinIO
            try:
                audio_object = minioСlient.get_object('interview-transcription-bucket', minio_filename)
                audio_content = io.BytesIO(audio_object.read())
            finally:
                audio_object.close()
                audio_object.release_conn()

            # Calculate the hash of the audio content
            hash_sha256 = hashlib.sha256(audio_content.getvalue()).hexdigest()
            # Create a response with the audio content and the hash
            response = make_response(send_file(audio_content, mimetype='audio/mpeg'))
            response.headers['X-File-Hash'] = hash_sha256
            
            return response
        else:
            return jsonify({'error': "Not found file"})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    
#Проверка загруженных файлов и куки
@app.route('/upload/getfiles', methods=['GET'])
def index():
    user_id = request.cookies.get('user_id')

    #Создание идентификатора куки, если куки нет
    if user_id is None:
        user_id = str(uuid.uuid4())
        # Устанавливаем куки с идентификатором для клиента
        response = make_response(jsonify({'user_id': user_id}))
        response.set_cookie('user_id', user_id, max_age=24*60*60*7)  
        return response
    else:   
        # Сформируем префикс для поиска файлов в MinIO
        prefix = f"{user_id}%mp%"
        # Получаем список всех объектов (файлов) в бакете MinIO
        objects = minioСlient.list_objects('interview-transcription-bucket', prefix=prefix, recursive=False)
        # Извлекаем имена файлов (filename) из списка объектов
        filenames = [obj.object_name.split('%')[-1] for obj in objects]
        return jsonify({'filenames': filenames})